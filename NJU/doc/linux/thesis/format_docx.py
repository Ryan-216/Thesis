"""将 pandoc docx 的样式修改为 njuthesis 研究生论文格式

格式规范来源: njuthesis.cls 1.4.3 + njuthesis-graduate.def
- 页面: A4, 上下2.54cm, 左右3.18cm
- 正文: 小四(12pt), 宋体, 1.5倍行距
- 一级标题(章): 小二(18pt), 黑体加粗, 居中
- 二级标题(节): 小三(15pt), 黑体加粗
- 三级标题: 四号(14pt), 黑体加粗
- 脚注: 小五(9pt)
- 表格: 五号(10.5pt)
- 页眉: small, 楷体
- 参考文献: 五号(10.5pt)
"""

import sys
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

def set_font_style(style, font_name_cn, font_name_en, size_pt, bold=False, italic=False):
    """设置样式的字体"""
    font = style.font
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.name = font_name_en
    # 设置东亚字体
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:eastAsia'), font_name_cn)
    rfonts.set(qn('w:ascii'), font_name_en)
    rfonts.set(qn('w:hAnsi'), font_name_en)

def set_para_spacing(style, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落间距和行距"""
    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def modify_styles(doc_path, out_path):
    doc = Document(doc_path)
    sd = {s.name: s for s in doc.styles}

    # === 页面设置 ===
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # === Normal 样式: 小四 宋体 1.5倍行距 ===
    style = sd.get('Normal')
    if style:
        set_font_style(style, '宋体', 'Times New Roman', 12)
        set_para_spacing(style, 1.5, 0, 0)

    # === Heading 1: 章标题 小二(18pt) 黑体加粗 居中 ===
    style = sd.get('Heading 1')
    if style:
        set_font_style(style, '黑体', 'Arial', 18, bold=True)
        set_para_spacing(style, 1.5, 12, 6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Heading 2: 节标题 小三(15pt) 黑体加粗 ===
    style = sd.get('Heading 2')
    if style:
        set_font_style(style, '黑体', 'Arial', 15, bold=True)
        set_para_spacing(style, 1.5, 8, 4)

    # === Heading 3: 四号(14pt) 黑体加粗 ===
    style = sd.get('Heading 3')
    if style:
        set_font_style(style, '黑体', 'Arial', 14, bold=True)
        set_para_spacing(style, 1.5, 6, 3)

    # === 表格样式 ===
    style = sd.get('Table')
    if style:
        set_font_style(style, '宋体', 'Times New Roman', 10.5)
        set_para_spacing(style, 1.25, 0, 0)

    doc.save(out_path)
    print(f"Saved: {out_path}")

def clean_crossrefs(doc):
    """清理交叉引用标签 [tab:xxx], [fig:xxx], [eq:xxx]
    跨 run 合并处理——将匹配段落的全文替换写入第一个 run，其余清空"""
    import re
    patterns = [
        (re.compile(r'\[tab:([^\]]+)\]'), r'【表\1】'),
        (re.compile(r'\[fig:([^\]]+)\]'), r'【图\1】'),
        (re.compile(r'\[eq:([^\]]+)\]'),  r'【公式\1】'),
    ]
    for para in doc.paragraphs:
        changed = False
        for pat, repl in patterns:
            if pat.search(para.text):
                changed = True
                break
        if changed:
            full = para.text
            for pat, repl in patterns:
                full = pat.sub(repl, full)
            # 遍历段落下所有 w:t 节点（包括 hyperlink/fldChar 内的嵌套 run）
            text_nodes = para._element.findall('.//' + qn('w:t'))
            if text_nodes:
                text_nodes[0].text = full
                for node in text_nodes[1:]:
                    node.text = ''

def clean_duplicate_citations(doc):
    """删除重复引用：）(FullCitation) 紧随 （Author Year） 之后且同年者，移除后者；
    也处理 Author（Year）... (Author Year) 模式"""
    import re
    # 模式1: （Author et al., Year）(FullCitation Year) 同括号内同年
    dup1 = re.compile(r'（([^）]*?(\d{4}))）\s*\(([^)]*\2[^)]*)\)')
    # 模式2: Author（Year）中间文本 (Author Year)
    dup2 = re.compile(r'([A-Z][a-z]+)（(\d{4})）([^（]*?)\(\1[^)]*\2[^)]*\)')
    for para in doc.paragraphs:
        full = para.text
        cleaned = full
        if dup1.search(cleaned):
            cleaned = dup1.sub(r'（\1）', cleaned)
        if dup2.search(cleaned):
            cleaned = dup2.sub(r'\1（\2）\3', cleaned)
        if cleaned != full:
            text_nodes = para._element.findall('.//' + qn('w:t'))
            if text_nodes:
                text_nodes[0].text = cleaned
                for node in text_nodes[1:]:
                    node.text = ''

def remove_cover_image(doc):
    """移除封面图片段落"""
    to_remove = []
    for i, para in enumerate(doc.paragraphs[:5]):
        xml = para._element.xml
        if 'wp:inline' in xml or 'wp:anchor' in xml or 'w:drawing' in xml:
            to_remove.append(para)
        elif para.text.strip() == '' and len(para.runs) > 0:
            for run in para.runs:
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    to_remove.append(para)
                    break
    for para in to_remove:
        p_elem = para._element
        p_elem.getparent().remove(p_elem)
    print(f"Removed {len(to_remove)} cover-related paragraphs")

if __name__ == '__main__':
    inp = sys.argv[1] if len(sys.argv) > 1 else '/tmp/thesis_v2_cover.docx'
    out = sys.argv[2] if len(sys.argv) > 2 else '/tmp/thesis_v3_formatted.docx'
    remove_cover = '--no-cover' in sys.argv
    clean_refs = '--clean-refs' in sys.argv

    doc = Document(inp)
    sd = {s.name: s for s in doc.styles}

    if remove_cover:
        remove_cover_image(doc)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = sd.get('Normal')
    if style:
        set_font_style(style, '\u5b8b\u4f53', 'Times New Roman', 12)
        set_para_spacing(style, 1.5, 0, 0)

    style = sd.get('Heading 1')
    if style:
        set_font_style(style, '\u9ed1\u4f53', 'Arial', 18, bold=True)
        set_para_spacing(style, 1.5, 12, 6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style = sd.get('Heading 2')
    if style:
        set_font_style(style, '\u9ed1\u4f53', 'Arial', 15, bold=True)
        set_para_spacing(style, 1.5, 8, 4)

    style = sd.get('Heading 3')
    if style:
        set_font_style(style, '\u9ed1\u4f53', 'Arial', 14, bold=True)
        set_para_spacing(style, 1.5, 6, 3)

    style = sd.get('Table')
    if style:
        set_font_style(style, '\u5b8b\u4f53', 'Times New Roman', 10.5)
        set_para_spacing(style, 1.25, 0, 0)

    if clean_refs:
        clean_crossrefs(doc)
        clean_duplicate_citations(doc)

    doc.save(out)
    print(f"Saved: {out}")
