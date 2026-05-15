"""插入封面图片到 pandoc 生成的 docx 前面"""
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = sys.argv[1] if len(sys.argv) > 1 else '/tmp/thesis_v2.docx'
out_path = sys.argv[2] if len(sys.argv) > 2 else '/tmp/thesis_v2_cover.docx'

doc = Document(doc_path)
first_para = doc.paragraphs[0]

# 插入 cover_1.png (封面)
p = first_para.insert_paragraph_before()
run = p.add_run()
run.add_picture('/tmp/cover_1.png', width=Inches(5.8))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 分页符
p2 = first_para.insert_paragraph_before()
p2.add_run().add_break()

doc.save(out_path)
print(f'Saved to {out_path}')
